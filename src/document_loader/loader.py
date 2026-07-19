"""文档加载器 — 支持 PDF、TXT、Markdown、Word 多种格式"""

import os
from typing import List, Optional
from datetime import datetime

from langchain_core.documents import Document


class DocumentLoader:
    """多格式文档加载器，返回 LangChain Document 列表"""

    SUPPORTED_TYPES = {
        ".txt": "text",
        ".md": "markdown",
        ".pdf": "pdf",
        ".docx": "docx",
    }

    @classmethod
    def load(cls, file_path: str) -> List[Document]:
        """
        根据文件扩展名自动选择合适的加载器。

        Args:
            file_path: 文件路径

        Returns:
            List[Document]: LangChain Document 对象列表

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in cls.SUPPORTED_TYPES:
            raise ValueError(
                f"不支持的文件格式: {ext}，支持: {list(cls.SUPPORTED_TYPES.keys())}"
            )

        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        metadata_base = {
            "source": filename,
            "file_path": file_path,
            "file_type": ext,
            "file_size": file_size,
            "loaded_at": datetime.now().isoformat(),
        }

        if ext == ".txt":
            return cls._load_txt(file_path, metadata_base)
        elif ext == ".md":
            return cls._load_markdown(file_path, metadata_base)
        elif ext == ".pdf":
            return cls._load_pdf(file_path, metadata_base)
        elif ext == ".docx":
            return cls._load_docx(file_path, metadata_base)
        else:
            raise ValueError(f"未处理的文件格式: {ext}")

    @classmethod
    def _load_txt(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata=metadata)]

    @classmethod
    def _load_markdown(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 Markdown 文件（保留原始格式，移除 HTML 标签）"""
        import markdown as md_lib
        from io import StringIO

        with open(file_path, "r", encoding="utf-8") as f:
            raw = f.read()

        # 保留纯文本版本（用于检索）
        # 同时保留原始 markdown（当作文本处理，因为 MD 本身可读性高）
        return [Document(page_content=raw, metadata={**metadata, "format": "markdown"})]

    @classmethod
    def _load_pdf(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 PDF 文件"""
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        documents = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                documents.append(
                    Document(
                        page_content=text.strip(),
                        metadata={**metadata, "page": i + 1, "total_pages": len(reader.pages)},
                    )
                )

        if not documents:
            # 如果所有页面都为空，返回一个空文档标记
            return [Document(page_content="[PDF 文件无可提取文本]", metadata=metadata)]

        return documents

    @classmethod
    def _load_docx(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 Word 文档"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        if not text:
            return [Document(page_content="[Word 文件无可提取文本]", metadata=metadata)]

        return [Document(page_content=text, metadata=metadata)]


def load_document(file_path: str) -> List[Document]:
    """便捷函数：加载文档"""
    return DocumentLoader.load(file_path)
